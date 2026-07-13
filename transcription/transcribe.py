"""
Script de transcrição de vídeo/áudio para texto.
Utiliza Whisper AI com detecção automática de device (CUDA/MPS/CPU).

Uso:
    python transcribe.py --input video.mp4
    python transcribe.py --input audio.mp3 --lang en --model small --output resultado.docx
"""

import argparse
import os
import sys
import tempfile
import shutil

import torch
import whisper
from docx import Document
from moviepy import VideoFileClip
from tqdm import tqdm

# Imports do projeto
from shared.utils.file_detector import get_file_type
from shared.config import (
    PROMPTS,
    WHISPER_MODELS,
    DEFAULT_MODEL,
    TRANSCRIPTION_CONFIG,
    TRANSCRIPTION_OUTPUT_DIR,
    AUDIO_CODEC,
    AUDIO_SAMPLE_RATE,
    AUDIO_CHANNELS,
)


# ── Detecção de device ──────────────────────────────────────────────────────

def get_device() -> str:
    """Seleciona automaticamente o melhor device disponível: CUDA > MPS > CPU."""
    if torch.cuda.is_available():
        return "cuda"
    if torch.backends.mps.is_available():
        return "mps"
    return "cpu"


# ── Extração de áudio ───────────────────────────────────────────────────────

def extract_audio(input_path: str, audio_path: str) -> None:
    """
    Extrai o áudio do vídeo (ou copia se já for áudio) para WAV mono 16 kHz.
    Usa context manager para garantir liberação de recursos.
    """
    file_type = get_file_type(input_path)

    if file_type == "video":
        print("🎞️  Extraindo áudio do vídeo...")
        with VideoFileClip(input_path) as clip:
            clip.audio.write_audiofile(
                audio_path,
                codec=AUDIO_CODEC,
                ffmpeg_params=[
                    "-ar", str(AUDIO_SAMPLE_RATE),
                    "-ac", str(AUDIO_CHANNELS)
                ],
                logger=None,  # suprime logs do moviepy
            )
    else:
        # Arquivo de áudio: copia diretamente para o temp
        print("🎵  Copiando arquivo de áudio...")
        shutil.copy2(input_path, audio_path)


# ── Transcrição ─────────────────────────────────────────────────────────────

def transcribe(audio_path: str, lang: str, model_name: str, device: str) -> list[str]:
    """
    Transcreve o áudio com Whisper e retorna lista de parágrafos.

    Melhorias de performance aplicadas:
    - device explícito (cuda/mps/cpu)
    - beam_size=1 (greedy decode): ~2x mais rápido que beam_size=5
    - condition_on_previous_text=False: evita repetições em vídeos longos
    - fp16=True em GPU, False em CPU (evita warning e instabilidade)
    """
    print(f"🧠  Transcrevendo com modelo '{model_name}' no device '{device}'...")

    model = whisper.load_model(model_name, device=device)

    result = model.transcribe(
        audio_path,
        language=lang,
        initial_prompt=PROMPTS.get(lang, ""),
        fp16=(device != "cpu"),
        **TRANSCRIPTION_CONFIG
    )

    paragraphs = []
    for seg in tqdm(result["segments"], desc="📝  Formatando"):
        text = seg["text"].strip()
        if text:
            paragraphs.append(text)

    return paragraphs


# ── Salvar DOCX ─────────────────────────────────────────────────────────────

def save_docx(paragraphs: list[str], output_file: str) -> None:
    """Salva os parágrafos transcritos em um arquivo DOCX."""
    # Garante que o diretório de output existe
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    doc.add_heading("Transcrição Completa", level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_file)


# ── CLI ─────────────────────────────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Transcreve vídeo ou áudio para DOCX usando Whisper AI.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Caminho para o arquivo de vídeo ou áudio de entrada.",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help=f"Caminho para o arquivo DOCX de saída (padrão: {TRANSCRIPTION_OUTPUT_DIR}/transcricao_<lang>.docx).",
    )
    parser.add_argument(
        "--lang", "-l",
        default="pt",
        choices=list(PROMPTS.keys()),
        help="Idioma da transcrição.",
    )
    parser.add_argument(
        "--model", "-m",
        default=DEFAULT_MODEL,
        choices=WHISPER_MODELS,
        help="Tamanho do modelo Whisper. Modelos menores são mais rápidos.",
    )
    return parser.parse_args()


# ── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    args = parse_args()

    # Define output file com caminho no diretório correto
    if args.output:
        output_file = args.output
    else:
        output_file = os.path.join(TRANSCRIPTION_OUTPUT_DIR, f"transcricao_{args.lang}.docx")
    
    device = get_device()

    print(f"🔊  Iniciando transcrição...")
    print(f"    Arquivo  : {args.input}")
    print(f"    Idioma   : {args.lang}")
    print(f"    Modelo   : {args.model}")
    print(f"    Device   : {device}")
    print(f"    Saída    : {output_file}")
    print()

    # Arquivo temporário gerenciado pelo SO — limpo automaticamente em caso de erro
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
        audio_path = tmp.name

    try:
        extract_audio(args.input, audio_path)
        paragraphs = transcribe(audio_path, args.lang, args.model, device)
        save_docx(paragraphs, output_file)
    except Exception as e:
        print(f"\n❌  Erro: {e}", file=sys.stderr)
        sys.exit(1)
    finally:
        if os.path.exists(audio_path):
            os.remove(audio_path)

    print(f"\n✅  Transcrição finalizada! Arquivo salvo como: {output_file}")


if __name__ == "__main__":
    main()
